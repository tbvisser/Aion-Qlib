"""Standalone DOCX report generator for contract review harness.

This script runs inside the Python sandbox. It reads workspace files from
/sandbox/workspace/ and produces a formatted DOCX report in /sandbox/output/.

Usage: The harness injects `sys.argv = ['script', '<output_filename>']` before
executing this script in the sandbox.
"""
import json
import os
import sys
from datetime import datetime

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ---------------------------------------------------------------------------
# Colors
# ---------------------------------------------------------------------------
RED_BG = "FDECEA"
RED_TEXT = "C0392B"
YELLOW_BG = "FFF9E6"
YELLOW_TEXT = "B8860B"
GREEN_BG = "E8F5E9"
GREEN_TEXT = "2E7D32"
HEADER_BG = "1B2A4A"
LIGHT_GRAY_BG = "F5F6FA"
ACCENT = "2E75B6"

RISK_COLORS = {
    "RED": {"bg": RED_BG, "text": RED_TEXT, "label": "CRITICAL (RED)"},
    "YELLOW": {"bg": YELLOW_BG, "text": YELLOW_TEXT, "label": "MODERATE (YELLOW)"},
    "GREEN": {"bg": GREEN_BG, "text": GREEN_TEXT, "label": "LOW (GREEN)"},
}


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def set_cell_shading(cell, color):
    """Set background shading on a table cell."""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)


def set_cell_text(cell, text, bold=False, size=10, color=None, italic=False):
    """Set text in a cell with formatting."""
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Arial"
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if italic:
        run.italic = True
    return run


def add_formatted_paragraph(doc, text, bold=False, size=11, color=None,
                            alignment=None, space_after=Pt(6), space_before=Pt(0)):
    """Add a paragraph with formatting."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Arial"
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if alignment:
        p.alignment = alignment
    p.paragraph_format.space_after = space_after
    p.paragraph_format.space_before = space_before
    return p


def load_json_file(path, default=None):
    """Load a JSON file, returning default on any error."""
    if default is None:
        default = {}
    try:
        with open(path, "r", encoding="utf-8") as f:
            return json.load(f)
    except (FileNotFoundError, json.JSONDecodeError):
        return default


def load_text_file(path, default=""):
    """Load a text file, returning default on error."""
    try:
        with open(path, "r", encoding="utf-8") as f:
            return f.read()
    except FileNotFoundError:
        return default


# ---------------------------------------------------------------------------
# Main report generation
# ---------------------------------------------------------------------------

def generate_report(output_filename):
    """Generate the DOCX report from workspace files."""
    workspace = "/sandbox/workspace"
    output_dir = "/sandbox/output"
    os.makedirs(output_dir, exist_ok=True)

    # Load workspace data
    classification = load_json_file(os.path.join(workspace, "classification.md"), {})
    risk_analysis = load_json_file(os.path.join(workspace, "risk-analysis.md"), [])
    redlines = load_json_file(os.path.join(workspace, "redlines.md"), [])
    report_text = load_text_file(os.path.join(workspace, "contract-review-report.md"))

    # Parse report as JSON if possible (it's the executive summary output)
    report_data = {}
    try:
        report_data = json.loads(report_text)
    except (json.JSONDecodeError, TypeError):
        pass

    # Extract key data
    contract_type = classification.get("contract_type", "Contract")
    parties = classification.get("parties", [])
    party_name = parties[0].get("name", "Report") if parties else "Report"

    overall_risk = report_data.get("overall_risk", "YELLOW").upper()
    recommendation = report_data.get("recommendation", "Review required")
    key_findings = report_data.get("key_findings", [])
    risk_breakdown = report_data.get("risk_breakdown", {})

    # Normalize risk analysis — could be list or dict with "assessments" key
    if isinstance(risk_analysis, dict):
        assessments = risk_analysis.get("assessments", [])
    elif isinstance(risk_analysis, list):
        assessments = risk_analysis
    else:
        assessments = []

    # Normalize redlines — could be list or dict with "redlines"/"assessments" key
    if isinstance(redlines, dict):
        redline_list = redlines.get("redlines", []) or redlines.get("assessments", [])
    elif isinstance(redlines, list):
        redline_list = redlines
    else:
        redline_list = []

    # Count risks
    risk_counts = {"RED": 0, "YELLOW": 0, "GREEN": 0}
    for a in assessments:
        level = (a.get("risk_level") or "").upper()
        if level in risk_counts:
            risk_counts[level] += 1

    # Use breakdown from report if available, else computed
    if not risk_breakdown:
        risk_breakdown = risk_counts

    # Create document
    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Arial"
    font.size = Pt(11)

    # -----------------------------------------------------------------------
    # TITLE PAGE
    # -----------------------------------------------------------------------
    # Spacer
    for _ in range(4):
        doc.add_paragraph()

    add_formatted_paragraph(
        doc, "CONFIDENTIAL", bold=True, size=12, color=RED_TEXT,
        alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(12),
    )

    add_formatted_paragraph(
        doc, f"{contract_type} Review", bold=True, size=26, color=HEADER_BG,
        alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(6),
    )

    add_formatted_paragraph(
        doc, "Contract Risk Assessment & Negotiation Workbook",
        size=14, color="666666",
        alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(24),
    )

    add_formatted_paragraph(
        doc, f"Prepared for: {party_name}", size=12, color="444444",
        alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(6),
    )

    add_formatted_paragraph(
        doc, f"Date: {datetime.now().strftime('%B %Y')}", size=12, color="444444",
        alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(36),
    )

    # Overall risk badge
    risk_color = RISK_COLORS.get(overall_risk, RISK_COLORS["YELLOW"])
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Overall Risk Rating:  ")
    run.font.size = Pt(14)
    run.font.name = "Arial"
    run.font.color.rgb = RGBColor.from_string("444444")
    badge_run = p.add_run(f" {overall_risk} ({risk_color['label'].split('(')[0].strip()}) ")
    badge_run.bold = True
    badge_run.font.size = Pt(14)
    badge_run.font.name = "Arial"
    badge_run.font.color.rgb = RGBColor.from_string("FFFFFF")
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{risk_color["text"]}" w:val="clear"/>')
    badge_run._r.get_or_add_rPr().append(shading)

    add_formatted_paragraph(
        doc, f"Recommendation: {recommendation}", bold=True, size=12,
        color=risk_color["text"],
        alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(0),
    )

    doc.add_page_break()

    # -----------------------------------------------------------------------
    # EXECUTIVE SUMMARY
    # -----------------------------------------------------------------------
    doc.add_heading("Executive Summary", level=1)

    # Summary text from the report
    detailed = report_data.get("detailed_report", "")
    if detailed:
        # Just use the first paragraph as the summary intro
        first_para = detailed.split("\n\n")[0] if "\n\n" in detailed else detailed[:500]
        add_formatted_paragraph(doc, first_para, size=11, space_after=Pt(12))

    # Risk breakdown cards table
    doc.add_heading("Risk Breakdown", level=2)

    card_table = doc.add_table(rows=1, cols=3)
    card_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for idx, level in enumerate(["RED", "YELLOW", "GREEN"]):
        cell = card_table.cell(0, idx)
        rc = RISK_COLORS[level]
        count = risk_breakdown.get(level, risk_breakdown.get(level.lower(), 0))
        set_cell_shading(cell, rc["bg"])
        cell.text = ""
        p1 = cell.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = p1.add_run(rc["label"])
        r1.bold = True
        r1.font.size = Pt(11)
        r1.font.name = "Arial"
        r1.font.color.rgb = RGBColor.from_string(rc["text"])
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = p2.add_run(f"{count} Clauses")
        r2.bold = True
        r2.font.size = Pt(18)
        r2.font.name = "Arial"
        r2.font.color.rgb = RGBColor.from_string(rc["text"])

    doc.add_paragraph()  # spacer

    # -----------------------------------------------------------------------
    # KEY FINDINGS
    # -----------------------------------------------------------------------
    doc.add_heading("Key Findings", level=2)

    if key_findings:
        for i, finding in enumerate(key_findings, 1):
            add_formatted_paragraph(doc, f"{i}. {finding}", size=11, space_after=Pt(6))
    else:
        add_formatted_paragraph(doc, "No key findings recorded.", size=11)

    doc.add_page_break()

    # -----------------------------------------------------------------------
    # DETAILED FINDINGS TABLE
    # -----------------------------------------------------------------------
    doc.add_heading("Detailed Findings & Proposed Redlines", level=1)

    if not redline_list:
        add_formatted_paragraph(
            doc, "No redlines required — all clauses acceptable.",
            size=11, space_after=Pt(12),
        )
    else:
        add_formatted_paragraph(
            doc,
            "The table below contains each reviewed clause, its risk classification, "
            "the original language, proposed revisions, and rationale for the change.",
            size=11, space_after=Pt(12),
        )

        # Header row
        table = doc.add_table(rows=1, cols=5)
        table.style = "Table Grid"
        headers = ["Clause Ref", "Risk", "Original Text", "Proposed Text", "Rationale"]
        header_widths = [Inches(0.8), Inches(0.6), Inches(2.2), Inches(2.2), Inches(1.7)]

        for idx, (text, width) in enumerate(zip(headers, header_widths)):
            cell = table.rows[0].cells[idx]
            set_cell_shading(cell, HEADER_BG)
            set_cell_text(cell, text, bold=True, size=9, color="FFFFFF")
            cell.width = width

        # Data rows
        for i, r in enumerate(redline_list):
            row = table.add_row()
            bg = LIGHT_GRAY_BG if i % 2 == 1 else None

            # Clause ref
            cell0 = row.cells[0]
            set_cell_text(cell0, r.get("clause_ref", "?"), bold=True, size=9)
            if bg:
                set_cell_shading(cell0, bg)

            # Risk level with color
            cell1 = row.cells[1]
            level = (r.get("risk_level") or "YELLOW").upper()
            rc = RISK_COLORS.get(level, RISK_COLORS["YELLOW"])
            set_cell_shading(cell1, rc["bg"])
            set_cell_text(cell1, level, bold=True, size=9, color=rc["text"])

            # Original text
            cell2 = row.cells[2]
            set_cell_text(cell2, r.get("original_text", "N/A"), size=8)
            if bg:
                set_cell_shading(cell2, bg)

            # Proposed text
            cell3 = row.cells[3]
            set_cell_text(cell3, r.get("proposed_text", "N/A"), size=8, color=ACCENT)
            if bg:
                set_cell_shading(cell3, bg)

            # Rationale
            cell4 = row.cells[4]
            rationale = r.get("rationale", "")
            set_cell_text(cell4, rationale, size=8)
            if bg:
                set_cell_shading(cell4, bg)

    doc.add_page_break()

    # -----------------------------------------------------------------------
    # ACCEPTABLE CLAUSES
    # -----------------------------------------------------------------------
    doc.add_heading("Acceptable Clauses (GREEN)", level=1)

    green_clauses = [
        a for a in assessments
        if (a.get("risk_level") or "").upper() == "GREEN"
    ]

    if green_clauses:
        refs = ", ".join(a.get("clause_ref", "?") for a in green_clauses)
        add_formatted_paragraph(
            doc,
            f"The following clauses were reviewed and found to be within standard "
            f"market positions: {refs}. No changes are recommended for these clauses.",
            size=11, space_after=Pt(12),
        )
    else:
        add_formatted_paragraph(
            doc, "No clauses were assessed as GREEN.", size=11, space_after=Pt(12),
        )

    # -----------------------------------------------------------------------
    # RECOMMENDED NEXT STEPS
    # -----------------------------------------------------------------------
    doc.add_heading("Recommended Next Steps", level=1)

    red_count = risk_counts["RED"]
    yellow_count = risk_counts["YELLOW"]

    # Derive steps from actual analysis
    steps = []
    if red_count > 0:
        steps.append(f"Prioritize negotiation on the {red_count} RED clauses immediately.")
    if yellow_count > 0:
        steps.append(f"Address the {yellow_count} YELLOW clauses as part of the broader negotiation, using fallback positions.")

    # Check if data protection clauses are flagged
    data_flagged = any(
        (a.get("category") or "").lower() in ("data protection", "data location", "breach notification")
        and (a.get("risk_level") or "").upper() in ("RED", "YELLOW")
        for a in assessments
    )
    if data_flagged:
        steps.append("Prepare and submit a formal Data Processing Addendum (DPA) to address data protection concerns.")

    if red_count > 0:
        steps.append("Do not execute the agreement until all RED issues have been resolved or formally accepted with documented risk acknowledgment.")
    elif yellow_count > 0:
        steps.append("Review YELLOW items and confirm acceptable risk tolerance before execution.")
    else:
        steps.append("The agreement appears acceptable. Proceed with standard legal review and execution.")

    for i, step in enumerate(steps, 1):
        add_formatted_paragraph(doc, f"{i}. {step}", size=11, space_after=Pt(6))

    # -----------------------------------------------------------------------
    # HEADERS & FOOTERS
    # -----------------------------------------------------------------------
    for section in doc.sections:
        header = section.header
        hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        hp.text = ""
        hr = hp.add_run(f"{contract_type} Review | {party_name}")
        hr.font.size = Pt(8)
        hr.font.name = "Arial"
        hr.font.color.rgb = RGBColor.from_string("999999")
        hr2 = hp.add_run("    CONFIDENTIAL")
        hr2.font.size = Pt(8)
        hr2.font.name = "Arial"
        hr2.font.color.rgb = RGBColor.from_string(RED_TEXT)
        hr2.bold = True

        footer = section.footer
        fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fp.text = ""
        fr = fp.add_run("Contract Review Report — Generated by AI Agent Platform")
        fr.font.size = Pt(8)
        fr.font.name = "Arial"
        fr.font.color.rgb = RGBColor.from_string("999999")

    # Save
    output_path = os.path.join(output_dir, output_filename)
    doc.save(output_path)
    print(f"DOCX report saved: {output_path}")


if __name__ == "__main__":
    filename = sys.argv[1] if len(sys.argv) > 1 else "Contract_Review_Report.docx"
    generate_report(filename)
